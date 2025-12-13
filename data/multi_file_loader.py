"""
Multi-File Data Loader Module

Handles loading and extracting data from multiple file types:
- Excel (.xlsx, .xls)
- Word (.docx, .doc)
- PDF (.pdf)
"""

import pandas as pd
import os
from typing import List, Dict, Optional, Union
from pathlib import Path
import warnings

# Try to import optional dependencies
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    try:
        import pdfplumber
        HAS_PDF = True
        USE_PDFPLUMBER = True
    except ImportError:
        HAS_PDF = False
        USE_PDFPLUMBER = False


class MultiFileLoader:
    """
    Loads and extracts data from multiple file types.
    """
    
    def __init__(self):
        """Initialize the multi-file loader."""
        self.supported_extensions = {
            '.xlsx': 'excel',
            '.xls': 'excel',
            '.docx': 'word',
            '.doc': 'word',
            '.pdf': 'pdf'
        }
        
    def detect_file_type(self, file_path: str) -> Optional[str]:
        """
        Detect file type from extension.
        
        Parameters:
        -----------
        file_path : str
            Path to file
            
        Returns:
        --------
        str or None
            File type ('excel', 'word', 'pdf') or None if unsupported
        """
        ext = Path(file_path).suffix.lower()
        return self.supported_extensions.get(ext)
    
    def load_excel(self, file_path: str) -> pd.DataFrame:
        """
        Load data from Excel file.
        
        Parameters:
        -----------
        file_path : str
            Path to Excel file
            
        Returns:
        --------
        pd.DataFrame
            Loaded data
        """
        from .loader import DataLoader
        loader = DataLoader()
        return loader.load_data(file_path)
    
    def extract_from_word(self, file_path: str) -> Dict:
        """
        Extract data from Word document.
        
        Looks for:
        - Tables (converts to DataFrames)
        - Key-value pairs in text
        - Financial data patterns
        
        Parameters:
        -----------
        file_path : str
            Path to Word file
            
        Returns:
        --------
        Dict
            Dictionary with extracted data and metadata
        """
        if not HAS_DOCX:
            raise ImportError(
                "python-docx not installed. Install with: pip install python-docx"
            )
        
        doc = Document(file_path)
        extracted_data = {
            'tables': [],
            'text': [],
            'key_values': {},
            'metadata': {}
        }
        
        # Extract all text
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text.strip())
        extracted_data['text'] = full_text
        
        # Extract tables
        for i, table in enumerate(doc.tables):
            table_data = []
            headers = None
            
            for row_idx, row in enumerate(table.rows):
                row_data = [cell.text.strip() for cell in row.cells]
                
                if row_idx == 0:
                    headers = row_data
                else:
                    if headers:
                        table_dict = dict(zip(headers, row_data))
                        table_data.append(table_dict)
            
            if table_data:
                df = pd.DataFrame(table_data)
                extracted_data['tables'].append({
                    'index': i,
                    'dataframe': df,
                    'headers': headers
                })
        
        # Try to extract key-value pairs from text
        extracted_data['key_values'] = self._extract_key_values(full_text)
        
        return extracted_data
    
    def extract_from_pdf(self, file_path: str) -> Dict:
        """
        Extract data from PDF file.
        
        Looks for:
        - Tables (converts to DataFrames)
        - Text content
        - Financial data patterns
        
        Parameters:
        -----------
        file_path : str
            Path to PDF file
            
        Returns:
        --------
        Dict
            Dictionary with extracted data and metadata
        """
        if not HAS_PDF:
            raise ImportError(
                "PDF library not installed. Install with: pip install PyPDF2 or pip install pdfplumber"
            )
        
        extracted_data = {
            'tables': [],
            'text': [],
            'key_values': {},
            'metadata': {}
        }
        
        if USE_PDFPLUMBER:
            # Use pdfplumber (better for tables)
            import pdfplumber
            
            with pdfplumber.open(file_path) as pdf:
                full_text = []
                
                for page in pdf.pages:
                    # Extract text
                    page_text = page.extract_text()
                    if page_text:
                        full_text.append(page_text)
                    
                    # Extract tables
                    tables = page.extract_tables()
                    for table in tables:
                        if table and len(table) > 1:
                            # First row as headers
                            headers = table[0]
                            data = table[1:]
                            
                            # Create DataFrame
                            try:
                                df = pd.DataFrame(data, columns=headers)
                                extracted_data['tables'].append({
                                    'dataframe': df,
                                    'headers': headers
                                })
                            except:
                                pass
                
                extracted_data['text'] = full_text
        else:
            # Use PyPDF2 (basic text extraction)
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                full_text = []
                
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        full_text.append(text)
                
                extracted_data['text'] = full_text
        
        # Try to extract key-value pairs
        all_text = ' '.join(extracted_data['text'])
        extracted_data['key_values'] = self._extract_key_values([all_text])
        
        return extracted_data
    
    def _extract_key_values(self, text_lines: List[str]) -> Dict:
        """
        Extract key-value pairs from text.
        
        Looks for patterns like:
        - "WACC: 8%"
        - "Investment: $20,000,000"
        - "Year 1: 1000"
        
        Parameters:
        -----------
        text_lines : List[str]
            List of text lines
            
        Returns:
        --------
        Dict
            Dictionary of extracted key-value pairs
        """
        import re
        
        key_values = {}
        all_text = ' '.join(text_lines)
        
        # Pattern 1: "Key: Value"
        pattern1 = r'([A-Za-z\s]+):\s*([\d,\.\$%]+)'
        matches = re.findall(pattern1, all_text)
        for key, value in matches:
            key = key.strip()
            # Clean value
            value = value.replace(',', '').replace('$', '').replace('%', '')
            try:
                if '.' in value:
                    key_values[key] = float(value)
                else:
                    key_values[key] = int(value)
            except:
                key_values[key] = value.strip()
        
        # Pattern 2: Financial terms
        financial_patterns = {
            r'WACC[:\s]+([\d\.]+)%?': 'wacc',
            r'Investment[:\s]+\$?([\d,]+)': 'rubicon_investment_total',
            r'Streaming[:\s]+([\d\.]+)%?': 'streaming_percentage',
            r'Tenor[:\s]+(\d+)': 'investment_tenor',
        }
        
        for pattern, key in financial_patterns.items():
            match = re.search(pattern, all_text, re.IGNORECASE)
            if match:
                value = match.group(1).replace(',', '')
                try:
                    if key == 'wacc' or 'percentage' in key:
                        key_values[key] = float(value) / 100
                    else:
                        key_values[key] = float(value)
                except:
                    pass
        
        return key_values
    
    def find_data_table(self, extracted_data: Dict) -> Optional[pd.DataFrame]:
        """
        Find the most likely data table from extracted data.
        
        Looks for tables with:
        - Year column
        - Financial data columns
        - Appropriate number of rows (around 20 for 20-year model)
        
        Parameters:
        -----------
        extracted_data : Dict
            Extracted data from file
            
        Returns:
        --------
        pd.DataFrame or None
            Best matching data table
        """
        if not extracted_data.get('tables'):
            return None
        
        best_table = None
        best_score = 0
        
        for table_info in extracted_data['tables']:
            df = table_info['dataframe']
            score = 0
            
            # Check for year column
            year_cols = [col for col in df.columns if 'year' in str(col).lower()]
            if year_cols:
                score += 10
            
            # Check for financial columns
            financial_keywords = ['price', 'cost', 'credit', 'revenue', 'cash', 'flow']
            for col in df.columns:
                col_lower = str(col).lower()
                if any(keyword in col_lower for keyword in financial_keywords):
                    score += 5
            
            # Check row count (prefer ~20 rows)
            row_count = len(df)
            if 15 <= row_count <= 25:
                score += 10
            elif 10 <= row_count <= 30:
                score += 5
            
            # Check for numeric data
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) >= 2:
                score += 5
            
            if score > best_score:
                best_score = score
                best_table = df
        
        return best_table
    
    def load_multiple_files(self, file_paths: List[str]) -> Dict:
        """
        Load and extract data from multiple files.
        
        Parameters:
        -----------
        file_paths : List[str]
            List of file paths
            
        Returns:
        --------
        Dict
            Dictionary with:
            - 'combined_data': pd.DataFrame (main data table)
            - 'assumptions': Dict (extracted assumptions)
            - 'sources': Dict (data from each file)
            - 'metadata': Dict (file information)
        """
        results = {
            'combined_data': None,
            'assumptions': {},
            'sources': {},
            'metadata': {}
        }
        
        all_tables = []
        all_assumptions = {}
        
        for file_path in file_paths:
            if not os.path.exists(file_path):
                warnings.warn(f"File not found: {file_path}")
                continue
            
            file_type = self.detect_file_type(file_path)
            if not file_type:
                warnings.warn(f"Unsupported file type: {file_path}")
                continue
            
            file_name = Path(file_path).name
            results['metadata'][file_name] = {
                'path': file_path,
                'type': file_type
            }
            
            try:
                if file_type == 'excel':
                    # Load Excel using existing loader
                    data = self.load_excel(file_path)
                    results['sources'][file_name] = {
                        'type': 'excel',
                        'data': data,
                        'assumptions': {}
                    }
                    all_tables.append(data)
                    
                elif file_type == 'word':
                    extracted = self.extract_from_word(file_path)
                    results['sources'][file_name] = extracted
                    
                    # Find best table
                    table = self.find_data_table(extracted)
                    if table is not None:
                        all_tables.append(table)
                    
                    # Merge assumptions
                    all_assumptions.update(extracted.get('key_values', {}))
                    
                elif file_type == 'pdf':
                    extracted = self.extract_from_pdf(file_path)
                    results['sources'][file_name] = extracted
                    
                    # Find best table
                    table = self.find_data_table(extracted)
                    if table is not None:
                        all_tables.append(table)
                    
                    # Merge assumptions
                    all_assumptions.update(extracted.get('key_values', {}))
                    
            except Exception as e:
                warnings.warn(f"Error processing {file_name}: {str(e)}")
                results['sources'][file_name] = {'error': str(e)}
        
        # Combine tables (prefer Excel data, then largest table)
        if all_tables:
            # Prefer Excel tables first
            excel_tables = [t for t in all_tables if hasattr(t, 'index')]
            if excel_tables:
                results['combined_data'] = excel_tables[0]  # Use first Excel table
            else:
                # Use largest table
                results['combined_data'] = max(all_tables, key=len)
        
        results['assumptions'] = all_assumptions
        
        return results

